"""Build the 8570 reference PDF from the archived cyber.mil HTML.

Pipeline:
    archive HTML  -->  parse tables  -->  python-docx (exact layout)
                                      -->  (Word COM)  -->  pdf

Layout follows Jeff's final direction (2026-04-20 / 2026-04-21):
    - Heading "DoD 8570.01-M Approved Baseline Certifications" at top.
    - Baseline cert table first — ONE continuous 3-column table with
      alternating bold header rows ("IAT Level I/II/III", "IAM Level I/II/III",
      etc.) and data rows whose cells contain vertical cert lists.
      Matches the visual style of the original DoD page exactly.
    - Thin black grid lines on every cell.
    - Provider table after.
    - Notes (GSE/GISF removal, CySA+ rename, CASP+/SecurityX rename, our
      reproduction annotation) after tables.
    - Provenance, why-this-exists, reference block, compiled-by at the
      very end.
"""
import subprocess
import sys
import tempfile
from copy import copy
from datetime import date
from pathlib import Path

from bs4 import BeautifulSoup
from bs4.element import NavigableString
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor


# ---------------------------------------------------------------------------
# HTML parsing
# ---------------------------------------------------------------------------

def _header_texts_with_sup(row) -> list[tuple[str, str | None]]:
    """Extract header cell (main_text, sup_text) tuples.
    Preserves footnote superscript markers separately so we can render them
    as real superscripts in DOCX.
    """
    out: list[tuple[str, str | None]] = []
    for td in row.find_all(["td", "th"]):
        td_copy = copy(td)
        sup_marker = None
        sup = td_copy.find("sup")
        if sup:
            sup_marker = sup.get_text(" ", strip=True)
            sup.decompose()
        out.append((td_copy.get_text(" ", strip=True), sup_marker))
    return out


def _cell_cert_entries(td) -> list[tuple[str, bool]]:
    """Split a data cell into a list of (cert_text, is_red) tuples.

    The archived DoD page uses <font color="red"> to flag certs on the
    pending-removal list (e.g. HCISPP, CCSP in certain cells). We
    preserve that visual cue in the DOCX output."""
    entries: list[tuple[str, bool]] = []

    def walk(node, red_parent=False):
        """Walk cell children; emit (text, red) for each <br>-delimited segment."""
        # Detect red font wrapper
        red_here = red_parent
        if getattr(node, "name", None) == "font":
            color = node.get("color", "").lower()
            if color == "red":
                red_here = True
        # Flatten children's text into segments separated by <br>
        segments: list[tuple[str, bool]] = [("", red_here)]
        for child in getattr(node, "children", []):
            if getattr(child, "name", None) == "br":
                segments.append(("", red_here))
            elif isinstance(child, NavigableString):
                segments[-1] = (segments[-1][0] + str(child), segments[-1][1])
            else:
                # Recurse; append their segments, merging at boundaries
                inner = walk_collect(child, red_here)
                if not inner:
                    continue
                # First inner segment merges into current tail
                segments[-1] = (segments[-1][0] + inner[0][0], segments[-1][1] or inner[0][1])
                for seg in inner[1:]:
                    segments.append(seg)
        return segments

    def walk_collect(node, red_parent=False):
        return walk(node, red_parent)

    for seg_text, seg_red in walk(td):
        text = seg_text.strip()
        if text:
            entries.append((text, seg_red))
    return entries


def parse_baseline_table(table) -> list[tuple[list[tuple[str, str | None]], list[list[tuple[str, bool]]]]]:
    """Return a list of (header_cells, data_cells) pairs.
    Each pair is one section (IAT, IAM, etc.) — header_cells is a list of
    (text, sup_marker) tuples; data_cells is a list of per-column cert entry
    lists, where each entry is (text, is_red).
    """
    out = []
    rows = list(table.find_all("tr"))
    for i in range(0, len(rows), 2):
        if i + 1 >= len(rows):
            break
        headers = _header_texts_with_sup(rows[i])
        data = [_cell_cert_entries(td) for td in rows[i + 1].find_all(["td", "th"])]
        out.append((headers, data))
    return out


def parse_provider_table(table) -> list[tuple[str, str]]:
    """Return list of (provider, cert_name) tuples, preserving source order."""
    out: list[tuple[str, str]] = []
    rows = list(table.find_all("tr"))
    for tr in rows[1:]:  # skip header
        cells = [td.get_text(" ", strip=True).replace("\n", " ")
                 for td in tr.find_all(["td", "th"])]
        if len(cells) >= 2:
            out.append((cells[0], cells[1]))
    return out


def parse_footnotes(soup: BeautifulSoup) -> list[str]:
    """Return DoD's substantive footnotes from the source page:
    - "1." and "2." numbered notes referenced by superscripts in the
      section headers (CSSP category rename, CCNA-Security rebrand);
    - "The GIAC GSE..." (GSE/GISF removal);
    - "** CySA+..." (CySA+/CSA+ rename).

    Drops the generic "* This organization is the sole propriety owner..."
    boilerplate — adds no value for this reference."""
    keys = ("1. ", "2. ", "The GIAC GSE", "** CySA+")
    out: list[str] = []
    for p in soup.find_all("p"):
        text = p.get_text(" ", strip=True)
        if any(text.startswith(k) for k in keys):
            out.append(text)
    return out


# ---------------------------------------------------------------------------
# python-docx construction
# ---------------------------------------------------------------------------

THIN_BLACK_BORDER_XML = """
<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
</w:tcBorders>
"""

HEADER_SHADING_XML = """
<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  w:val="clear" w:color="auto" w:fill="F2F2F2"/>
"""

def _set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    # Remove existing tcBorders if any
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    from lxml import etree
    tc_pr.append(etree.fromstring(THIN_BLACK_BORDER_XML))


def _set_cell_header_shading(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    from lxml import etree
    tc_pr.append(etree.fromstring(HEADER_SHADING_XML))


def _write_section_header_row(row, headers: list[tuple[str, str | None]]) -> None:
    for i, (text, sup_marker) in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        # Preserve the superscript footnote markers (1, 2) from the source
        # DoD page — they link to the numbered footnotes in the Notes block.
        if sup_marker:
            sup_run = p.add_run(sup_marker)
            sup_run.font.superscript = True
            sup_run.font.size = Pt(9)
            sup_run.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_borders(cell)
        _set_cell_header_shading(cell)


def _write_data_row(row, data_cells: list[list[tuple[str, bool]]],
                    ncols: int) -> None:
    for i in range(ncols):
        cell = row.cells[i]
        cell.text = ""
        entries = data_cells[i] if i < len(data_cells) else []
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _set_cell_borders(cell)
        if not entries:
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            continue
        # First entry goes in the existing paragraph; subsequent entries
        # get new paragraphs so each cert is on its own line.
        for j, (text, is_red) in enumerate(entries):
            if j == 0:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(text)
            run.font.size = Pt(9)
            # Preserve the red-font markers from the original DoD page
            # (certs that were recent additions at time of publication).
            # Explained in the Notes section below the matrix.
            if is_red:
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def add_baseline_table(doc, blocks) -> None:
    """Add a single 3-column table covering all section pairs (header row +
    data row × 5). Each cell gets grid-line borders; header rows are shaded
    light gray."""
    total_rows = len(blocks) * 2
    table = doc.add_table(rows=total_rows, cols=3)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Set table grid layout style (ensures borders honor our per-cell setting)
    tbl_pr = table._element.tblPr
    from lxml import etree
    borders_xml = """<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
    </w:tblBorders>"""
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(etree.fromstring(borders_xml))

    row_idx = 0
    for headers, data_cells in blocks:
        # Pad headers/data to 3 columns
        while len(headers) < 3:
            headers.append(("", None))
        while len(data_cells) < 3:
            data_cells.append([])
        _write_section_header_row(table.rows[row_idx], headers[:3])
        _write_data_row(table.rows[row_idx + 1], data_cells[:3], 3)
        row_idx += 2


def add_provider_table(doc, providers: list[tuple[str, str]]) -> None:
    """Add the 2-column providers table with grid lines."""
    table = doc.add_table(rows=len(providers) + 1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Table-level borders
    from lxml import etree
    tbl_pr = table._element.tblPr
    borders_xml = """<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
    </w:tblBorders>"""
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(etree.fromstring(borders_xml))

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(["Certification Provider", "Certification Name"]):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_borders(cell)
        _set_cell_header_shading(cell)
    # Data rows
    for i, (provider, cert) in enumerate(providers, start=1):
        row = table.rows[i]
        for col, text in enumerate([provider, cert]):
            cell = row.cells[col]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_borders(cell)


def build_docx(out_path: Path, baseline_blocks, provider_rows,
               footnotes: list[str]) -> None:
    compile_date = date.today().isoformat()
    doc = Document()

    # Page setup: Letter portrait, tight margins so everything fits on one page
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.page_height = Inches(11.0)
        section.page_width = Inches(8.5)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("DoD 8570.01-M Approved Baseline Certifications")
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Baseline table
    add_baseline_table(doc, baseline_blocks)

    def section_header(text: str, space_before: int = 4) -> None:
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(1)
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(10)

    def bullet(text: str) -> None:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run("• " + text)
        r.font.size = Pt(8)

    def body(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(text)
        r.font.size = Pt(8)

    # ---- Notes (original DoD footnotes + our annotations) ----
    section_header("Notes")
    bullet(
        "This document reproduces the DoD 8570 Approved Baseline Certifications list "
        "as DoD last published it (snapshot date below). Cert names, vendor branding, "
        "red-font markers, and superscript footnote references are preserved as "
        "originally published; some annotations may no longer be current."
    )
    bullet(
        "Red-font entries in the matrix (HCISPP, CCSP) were flagged by DoD as recent "
        "additions to the approved list at time of publication. That currency note is "
        "no longer meaningful in this 2024-snapshot reproduction."
    )
    bullet(
        "CASP+ / SecurityX: CompTIA renamed the CompTIA Advanced Security Practitioner "
        "(CASP+) certification to SecurityX on 17 December 2024, coinciding with the "
        "release of exam version CAS-005 (V5). Same credential; any reference to CASP+ "
        "should be read as what CompTIA now calls SecurityX."
    )
    for fn in footnotes:
        bullet(fn)

    # ---- Provenance ----
    section_header("Provenance")
    body(
        "Reproduced from a web.archive.org snapshot of "
        "public.cyber.mil/wid/cwmp/dod-approved-8570-baseline-certifications/ "
        "dated 2024-01-30 — the last publicly archived version of this page before DoD "
        "removed it following the DoDM 8140.03 transition."
    )
    body(
        "Archive URL: https://web.archive.org/web/20240130012654/"
        "https://public.cyber.mil/wid/cwmp/dod-approved-8570-baseline-certifications/"
    )

    # ---- Explicit page break so the Providers table starts on page 2 ----
    break_para = doc.add_paragraph()
    break_run = break_para.add_run()
    from docx.enum.text import WD_BREAK
    break_run.add_break(WD_BREAK.PAGE)

    # ---- IA Workforce Certification Providers (page 2) ----
    providers_hdr = doc.add_paragraph()
    providers_hdr.paragraph_format.space_before = Pt(0)
    providers_hdr.paragraph_format.space_after = Pt(4)
    r = providers_hdr.add_run("IA Workforce Certification Providers")
    r.bold = True
    r.font.size = Pt(12)
    add_provider_table(doc, provider_rows)

    # ---- Why this document exists (bottom of page 2) ----
    section_header("Why this document exists", space_before=10)
    body(
        "DoD 8570.01-M was superseded by DoDM 8140.03 in 2023. When public.cyber.mil "
        "removed the 8570 baseline page, the authoritative list that many still-active "
        "contracts reference by name lost its public home. This document preserves the "
        "list so contract officers, CORs, and compliance staff can still cite an "
        "authoritative record. It is a reference reproduction, not a policy document. "
        "For current cybersecurity workforce qualification requirements see DoDM 8140.03 "
        "and the DoD Cyber Workforce Qualifications Matrices at "
        "www.cyber.mil/dod-workforce-innovation-directorate/dod8140/qualification-matrices."
    )

    # ---- Compiled by — at the very end ----
    tail = doc.add_paragraph()
    tail.paragraph_format.space_before = Pt(10)
    tail.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = tail.add_run(
        f"Compiled by Jeff Krueger on {compile_date}. "
        "Snapshot date 2024-01-30. Not legal advice."
    )
    run.italic = True
    run.font.size = Pt(8)

    doc.save(out_path)


def word_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """Open DOCX in Word and export to PDF via ExportAsFixedFormat."""
    import win32com.client  # type: ignore[import-untyped]
    import win32com.client.dynamic

    # Ensure target directory exists and no stale PDF is locked
    pdf_path = Path(pdf_path).resolve()
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    word = win32com.client.dynamic.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0  # wdAlertsNone
    # Force-disable Protected View / macros so Word doesn't open the DOCX
    # read-only (which blocks ExportAsFixedFormat).
    try:
        word.AutomationSecurity = 3  # msoAutomationSecurityForceDisable
    except Exception:
        pass
    try:
        doc = word.Documents.Open(FileName=str(docx_path.resolve()),
                                  ConfirmConversions=False,
                                  ReadOnly=False,
                                  AddToRecentFiles=False,
                                  NoEncodingDialog=True)
        # If somehow opened read-only, flip it.
        try:
            if getattr(doc, "ReadOnly", False):
                doc.Application.ActiveDocument.ReadOnlyRecommended = False
        except Exception:
            pass
        try:
            # Word's PDF export — analogous to Excel's ExportAsFixedFormat.
            # ExportFormat = 17 = wdExportFormatPDF.
            doc.ExportAsFixedFormat(
                OutputFileName=str(pdf_path),
                ExportFormat=17,
                OpenAfterExport=False,
                OptimizeFor=0,    # wdExportOptimizeForPrint
                Range=0,          # wdExportAllDocument
                Item=0,           # wdExportDocumentContent
                IncludeDocProps=True,
                KeepIRM=False,
                CreateBookmarks=0,
                DocStructureTags=True,
                BitmapMissingFonts=True,
                UseISO19005_1=False,
            )
        finally:
            doc.Close(SaveChanges=0)
    finally:
        word.Quit()


def render_pdf(html_path: str | Path, out_pdf: str | Path) -> Path:
    html = Path(html_path).read_text(encoding="utf-8")
    soup = BeautifulSoup(html, "lxml")
    baseline_table = soup.find("table", id="tablepress-iawip-approved_baseline_certifications")
    provider_table = soup.find("table", id="tablepress-iawip-certification_providers")
    baseline_blocks = parse_baseline_table(baseline_table) if baseline_table else []
    provider_rows = parse_provider_table(provider_table) if provider_table else []
    footnotes = parse_footnotes(soup)
    # Write DOCX to a non-temp location (gitignored data/ folder) so Word's
    # Protected View doesn't open the file read-only — Office blocks PDF
    # export from files opened through Protected View.
    out_pdf = Path(out_pdf).resolve()
    intermediate_dir = Path("data")
    intermediate_dir.mkdir(exist_ok=True)
    docx_path = intermediate_dir / "8570-intermediate.docx"
    build_docx(docx_path, baseline_blocks, provider_rows, footnotes)
    word_docx_to_pdf(docx_path, out_pdf)
    return out_pdf


if __name__ == "__main__":
    html = Path(
        sys.argv[1]
        if len(sys.argv) > 1
        else "8570/sources/cyber-mil-snapshot-20240130.html"
    )
    out = Path(
        sys.argv[2] if len(sys.argv) > 2 else "8570/8570-baseline-reference.pdf"
    )
    render_pdf(html, out)
    print(f"wrote {out}")
